from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    document = Document()
    
    # helper for style
    def add_heading(text, level=1):
        h = document.add_heading(text, level=level)
        return h

    def add_para(text, bold=False):
        p = document.add_paragraph()
        run = p.add_run(text)
        if bold:
            run.bold = True
        return p

    # Title
    title = document.add_heading('Technical Implementation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = document.add_paragraph('EarthApp: Offline Mesh Communication System')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_page_break()
    
    # 1. Executive Summary
    add_heading('1. Executive Summary', level=1)
    document.add_paragraph(
        "EarthApp is a specialized Android application designed to facilitate offline communication "
        "during emergency situations where traditional infrastructure (Cellular, Internet) is unavailable. "
        "It leverages ad-hoc mesh networking technologies, primarily Wi-Fi Direct (P2P), to create a "
        "decentralized network of devices."
    )
    
    # 2. Key Technical Specifications
    add_heading('2. Key Technical Specifications', level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parameter'
    hdr_cells[1].text = 'Specification'
    
    specs = [
        ('Application Name', 'EarthApp'),
        ('Architecture', 'MVVM (Model-View-ViewModel)'),
        ('Min / Target SDK', 'API 24 (Android 7.0) / API 36 (Android 14)'),
        ('Database', 'Room Persistence Library (SQLite)'),
        ('Network Protocol', 'Wi-Fi Direct (P2P) + TCP Sockets'),
        ('Port Used', '8888'),
        ('Data Format', 'JSON (via Gson)'),
        ('Concurrency', 'Java Threads, Executors, Handlers')
    ]
    
    for item, val in specs:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = val
        
    # 3. System Architecture
    add_heading('3. System Architecture', level=1)
    document.add_paragraph(
        "The application follows the recommended Google App Architecture Guide, processing data flows "
        "unidirectionally where possible."
    )
    
    # 3.1 ViewModel
    add_heading('3.1 View Layer & State Management', level=2)
    document.add_paragraph(
        "The MainViewModel acts as the central state holder, surviving configuration changes. "
        "It exposes LiveData observables to the MainActivity:"
    )
    document.add_paragraph("- selfUser: Tracks the local user's registration state.", style='List Bullet')
    document.add_paragraph("- allMessages: Observes the database for new incoming chat/emergency messages.", style='List Bullet')
    document.add_paragraph("- locationHelper: Provides real-time GPS updates.", style='List Bullet')
    
    # 3.2 Repository
    add_heading('3.2 Repository Layer', level=2)
    document.add_paragraph(
        "The UserManager class serves as the Repository. It abstracts the data sources (Local DB and Mesh Network) "
        "from the UI. When a message is sent from the UI, the Repository writes it to the local database and "
        "simultaneously relays it to the MeshManager for broadcasting."
    )

    # 4. Mesh Networking Implementation
    add_heading('4. Detailed Mesh Networking Logic', level=1)
    document.add_paragraph(
        "The core innovation of EarthApp is its custom implementation of Wi-Fi Direct networking, found in the "
        "'com.example.earthapp.service' package."
    )
    
    add_heading('4.1 Discovery and Connection', level=2)
    document.add_paragraph(
        "The MeshManager uses 'WifiP2pManager.discoverPeers()' to scan for nearby devices. "
        "Upon finding a peer, it initiates a connection. The Android OS negotiates the Group Owner (GO) "
        "role. The app adapts its role based on this negotiation:"
    )
    
    document.add_paragraph(
        "1. Host (Group Owner): The device starts a 'MeshServer' instance listening on TCP Port 8888.\n"
        "2. Client: The device obtains the GO's IP address and starts a 'MeshClient' to connect to that IP on Port 8888."
    )
    
    add_heading('4.2 MeshServer Logic', level=2)
    document.add_paragraph(
        "The MeshServer uses a CachedThreadPool to handle multiple incoming client connections simultaneously. "
        "It maintains a synchronized list of active client sockets. Key behaviors:"
    )
    document.add_paragraph("- Message Relaying: When a message is received from one client, the Server broadcasts it to all other connected clients.", style='List Bullet')
    document.add_paragraph("- Persistence: All incoming messages are immediately saved to the local Room database via an async executor.", style='List Bullet')
    
    add_heading('4.3 Message Structure & Flooding', level=2)
    document.add_paragraph(
        "To ensure messages propagate beyond immediate neighbors, the 'Message' entity includes a 'TTL' (Time-To-Live) field. "
        "When a node receives a message with TTL > 0:"
    )
    document.add_paragraph("1. It decrements the TTL.", style='List Bullet')
    document.add_paragraph("2. It rebroadcasts the message to its own known peers.", style='List Bullet')
    document.add_paragraph("This basic flooding algorithm allows for multi-hop communication capability.", style='List Bullet')

    # 5. Data Model
    add_heading('5. Data Model Schema', level=1)
    document.add_paragraph(
        "Data persistence is handled by Room. The primary entity is 'Message'."
    )
    
    table_db = document.add_table(rows=1, cols=3)
    table_db.style = 'Table Grid'
    hdr = table_db.rows[0].cells
    hdr[0].text = 'Field'
    hdr[1].text = 'Type'
    hdr[2].text = 'Description'
    
    schema = [
        ('messageId', 'String (PK)', 'Unique UUID for the message'),
        ('content', 'String', 'The payload text (e.g. "HELP")'),
        ('type', 'Integer', '0=SAFE, 1=HELP, 2=CUSTOM'),
        ('latitude', 'Double', 'Sender\'s GPS Latitude'),
        ('longitude', 'Double', 'Sender\'s GPS Longitude'),
        ('ttl', 'Integer', 'Hop limit counter')
    ]
    
    for field, type_, desc in schema:
        row = table_db.add_row().cells
        row[0].text = field
        row[1].text = type_
        row[2].text = desc

    # 6. User Interface & Logic
    add_heading('6. UI/UX Logic', level=1)
    document.add_paragraph(
        "The current UI is functional and prioritizes speed in emergencies."
    )
    add_heading('6.1 Message Display', level=2)
    document.add_paragraph(
        "The MessageAdapter dynamically colors messages to alert users:"
    )
    document.add_paragraph("- RED: Emergency (Type 1)", style='List Bullet')
    document.add_paragraph("- GREEN: Safety Status (Type 0)", style='List Bullet')
    document.add_paragraph("- BLUE: Chat Messages", style='List Bullet')
    
    add_heading('6.2 Distance Calculation', level=2)
    document.add_paragraph(
        "The app calculates the distance to the message sender using 'android.location.Location.distanceBetween()'. "
        "This distance is displayed in meters next to the sender's name, aiding in physical rescue efforts."
    )

    # 7. Conclusion
    add_heading('7. Conclusion', level=1)
    document.add_paragraph(
        "EarthApp successfully implements a proof-of-concept for offline mesh networking. "
        "By combining Wi-Fi Direct for high-bandwidth local connections and a flooding algorithm for propagation, "
        "it provides a vital communication line when standard networks fail."
    )

    document.save('EarthApp_HighLevel_Technical_Report.docx')
    print("Report generated successfully: EarthApp_HighLevel_Technical_Report.docx")

if __name__ == "__main__":
    create_report()
